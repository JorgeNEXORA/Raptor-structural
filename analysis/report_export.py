from analysis.history import build_comparison_text
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
except Exception:  # pragma: no cover
    Document = None


class ReportExporter:
    def __init__(self):
        if Document is None:
            raise ImportError("python-docx não está instalado. Instala com: pip install python-docx")

    def _set_base_style(self, document):
        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    def _add_title_block(self, document, project):
        p = document.add_paragraph()
        r = p.add_run("RELATÓRIO ESTRUTURAL - MVP")
        r.bold = True
        r.font.size = Pt(16)

        document.add_paragraph(f"Projeto: {project.name}")
        document.add_paragraph(f"Local: {project.location}")
        document.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    def _add_summary(self, document, project):
        document.add_heading("1. Resumo", level=1)
        document.add_paragraph(f"Pilares: {len(project.columns)}")
        document.add_paragraph(f"Vigas: {len(project.beams)}")
        document.add_paragraph(f"Lajes: {len(project.slabs)}")
        document.add_paragraph(f"Sapatas: {len(project.footings)}")
        document.add_paragraph(f"Vigas de amarração: {len(getattr(project, 'tie_beams', []))}")

    def _add_slabs(self, document, project):
        document.add_heading("2. Lajes", level=1)
        table = document.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Span (m)"
        hdr[2].text = "Gk"
        hdr[3].text = "Qk"
        hdr[4].text = "ULS"
        hdr[5].text = "Msd"

        for s in project.slabs:
            row = table.add_row().cells
            row[0].text = s.id
            row[1].text = f"{s.span_m:.2f}"
            row[2].text = f"{s.gk_kn_m2:.2f}"
            row[3].text = f"{s.qk_kn_m2:.2f}"
            row[4].text = f"{s.result.sd_uls_kn_m2:.2f}"
            row[5].text = f"{s.result.msd_knm_m:.2f}"

    def _add_beams(self, document, project):
        document.add_heading("3. Vigas", level=1)
        table = document.add_table(rows=1, cols=7)
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Nós"
        hdr[2].text = "ULS"
        hdr[3].text = "Msd"
        hdr[4].text = "Vsd"
        hdr[5].text = "Armadura"
        hdr[6].text = "Estribos"

        for b in project.beams:
            row = table.add_row().cells
            row[0].text = b.id
            row[1].text = f"{b.start_node}->{b.end_node}"
            row[2].text = f"{b.result.sd_uls_kn_m:.2f}"
            row[3].text = f"{b.result.msd_knm:.2f}"
            row[4].text = f"{b.result.vsd_kn:.2f}"
            row[5].text = b.reinforcement_result.get("bottom_text", "-")
            row[6].text = b.reinforcement_result.get("stirrups_text", "-")

    def _add_footings(self, document, project):
        document.add_heading("4. Sapatas", level=1)
        table = document.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Tipo"
        hdr[2].text = "Nsd"
        hdr[3].text = "Sigma"
        hdr[4].text = "As ad"
        hdr[5].text = "Solo / Punçoamento"

        for f in project.footings:
            row = table.add_row().cells
            row[0].text = f.id
            row[1].text = f.footing_type.value
            row[2].text = f"{f.result.nsd_kn:.2f}"
            row[3].text = f"{f.result.soil_stress_mpa:.3f}"
            row[4].text = f"{f.result.adopted_as_cm2:.2f}"
            row[5].text = f"σmin={f.result.sigma_min_mpa:.3f}; σmax={f.result.sigma_max_mpa:.3f}; P={f.result.punching_vsd_kn:.2f}/{f.result.punching_vrd_kn:.2f}"

    def _add_ties(self, document, project):
        document.add_heading("5. Vigas de amarração/equilíbrio", level=1)
        if not getattr(project, "tie_beams", None):
            document.add_paragraph("Sem vigas de amarração/equilíbrio.")
            return

        table = document.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Ligação"
        hdr[2].text = "Vão"
        hdr[3].text = "T"
        hdr[4].text = "As req / adotar"

        for t in project.tie_beams:
            row = table.add_row().cells
            row[0].text = t.id
            row[1].text = f"{t.start_footing_id}->{t.end_footing_id}"
            row[2].text = f"{t.span_m:.2f}"
            row[3].text = f"{t.tie_force_kn:.2f}"
            row[4].text = f"{t.required_as_cm2:.2f} / {t.adopted_bars}"

    def _add_scores(self, document, project):
        document.add_heading("6. Score global", level=1)
        if not getattr(project, "project_scores", None):
            document.add_paragraph("Sem score calculado.")
            return
        for k, v in project.project_scores.items():
            document.add_paragraph(f"{k}: {v:.2f}")

    def _add_advice(self, document, project):
        document.add_heading("7. Modo Engenheiro", level=1)
        if not getattr(project, "advice_messages", None):
            document.add_paragraph("Sem recomendações.")
            return
        for m in project.advice_messages:
            document.add_paragraph(m)

    def _add_comparison(self, document, project):
        document.add_heading("8. Comparação antes/depois", level=1)
        snaps = getattr(project, "history_snapshots", []) or []
        if len(snaps) < 2:
            document.add_paragraph("Sem histórico suficiente para comparação.")
            return
        lines = build_comparison_text(snaps[-2], snaps[-1])
        for line in lines:
            document.add_paragraph(line)

    def _add_alerts(self, document, project):
        document.add_heading("9. Alertas", level=1)
        if not project.alerts:
            document.add_paragraph("Sem alertas.")
            return
        for a in project.alerts:
            document.add_paragraph(f"[{a.level}] {a.message}")

    def export_docx(self, project, output_path):
        document = Document()
        self._set_base_style(document)
        self._add_title_block(document, project)
        document.add_paragraph("")
        self._add_summary(document, project)
        self._add_slabs(document, project)
        self._add_beams(document, project)
        self._add_footings(document, project)
        self._add_ties(document, project)
        self._add_scores(document, project)
        self._add_advice(document, project)
        self._add_comparison(document, project)
        self._add_alerts(document, project)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        document.save(output_path)
        return output_path
